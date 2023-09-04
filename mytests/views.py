import json
import string

from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator
from django.http import JsonResponse
from django.shortcuts import redirect

from createtests.models import UserTest
from django.contrib import messages
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from django.shortcuts import render
import docx
from django.http import HttpResponse
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, TableStyle
from reportlab.lib import colors
from reportlab.platypus import Table


# ----------------------txt view---------------
from payment.models import UserMembership


@login_required(login_url='/authentication/login')
def download_student_view_txt(request, id):
    user_tests = UserTest.objects.get(owner=request.user, pk=id)
    alphabet = string.ascii_uppercase

    # Function to create the text file for teacher view
    def create_student_view_txt(user_tests):
        content = f"{user_tests.header}\n\n{user_tests.subtitle}\n\n{user_tests.institution}\n\n{user_tests.add_header_info}\n\n"

        for question_id, question_data in user_tests.questions.items():
            content += f"Question {question_id[1:]}:\n{question_data['question']}\n"

            if 'answers' in question_data:
                for index, answer in enumerate(question_data['answers'], start=1):
                    content += f"{alphabet[index-1]}. {answer}\n"

            content += "\n"  # Adding a blank line between questions

        content += "Grading\nPercentage\tGrade\n"
        for grade_info in user_tests.grades:
            content += f"{grade_info['percentage']}%\t\t{grade_info['grade']}\n"

        content += f"\n{user_tests.footer}"

        return content

    # Create the text file content
    content = create_student_view_txt(user_tests)

    # Create a response with the text content as a downloadable file
    response = HttpResponse(content, content_type='text/plain')
    response['Content-Disposition'] = f'attachment; filename={user_tests.header}_student_view.txt'

    return response


@login_required(login_url='/authentication/login')
def download_teacher_view_txt(request, id):
    user_tests =UserTest.objects.get(owner=request.user, pk=id)
    alphabet = string.ascii_uppercase

    # Function to create the text file for teacher view
    def create_teacher_view_txt(user_tests):
        content = f"{user_tests.header}\n\n{user_tests.subtitle}\n\n{user_tests.institution}\n\n{user_tests.add_header_info}\n\n"

        for question_id, question_data in user_tests.questions.items():
            content += f"Question {question_id[1:]}:\n{question_data['question']}\n"

            if 'answers' in question_data:
                for index, answer in enumerate(question_data['answers'], start=1):
                    content += f"{alphabet[index-1]}. {answer}\n"
            content += "\n"  # Adding a blank line between questions

            if 'correct_answer' in question_data:
                content += f"Correct Answer: {', '.join(str(alphabet[answer-1]) for answer in question_data['correct_answer'])}\n"

            if question_data['explanation']:
                content += f"Explanation: {question_data['explanation']}\n"

            content += "\n"  # Adding a blank line between questions

        content += "Grading\nPercentage\tGrade\n"
        for grade_info in user_tests.grades:
            content += f"{grade_info['percentage']}%\t\t{grade_info['grade']}\n"

        content += f"\n{user_tests.footer}"

        return content

    # Create the text file content
    content = create_teacher_view_txt(user_tests)

    # Create a response with the text content as a downloadable file
    response = HttpResponse(content, content_type='text/plain')
    response['Content-Disposition'] = f'attachment; filename={user_tests.header}_teacher_view.txt'

    return response


@login_required(login_url='/authentication/login')
def download_student_view_pdf(request, id):
    user_tests =UserTest.objects.get(owner=request.user, pk=id)
    alphabet = string.ascii_uppercase

    # Function to create the PDF document for teacher view
    def create_student_view_pdf(user_tests):
        response = HttpResponse(content_type='application/pdf')
        filename = f'{user_tests.header}_student_view.pdf'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'

        doc = SimpleDocTemplate(response, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []

        header_style = styles['Title']
        header_style.alignment = 1  # Center alignment
        elements.append(Paragraph(user_tests.header, header_style))

        subtitle_style = styles['Heading3']
        subtitle_style.alignment = 1  # Center alignment
        elements.append(Paragraph(user_tests.subtitle, subtitle_style))

        institution_style = styles['Heading3']
        institution_style.alignment = 1  # Center alignment
        elements.append(Paragraph(user_tests.institution, institution_style))

        add_header_info_style = styles['Heading3']
        add_header_info_style.alignment = 1  # Center alignment
        elements.append(Paragraph(user_tests.add_header_info, add_header_info_style))

        for question_id, question_data in user_tests.questions.items():
            elements.append(Paragraph(f"Question {question_id[1:]}:", styles['Heading2']))
            elements.append(Paragraph(question_data['question'], styles['Normal']))
            elements.append(Spacer(1, 12))

            if 'answers' in question_data:
                for index, answer in enumerate(question_data['answers'], start=1):
                    elements.append(Paragraph(f"{str(alphabet[index-1])}. {answer}", styles['Normal']))

        elements.append(Paragraph("Grades", styles['Heading2']))
        # Add table with grades
        table_data = [['Percentage', 'Grade']]
        for grade in user_tests.grades:
            table_data.append([f"{grade['percentage']}%", grade['grade']])

        table_style = TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                  ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                  ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                  ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                  ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                  ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                  ('GRID', (0, 0), (-1, -1), 1, colors.black)])

        # Create the table and apply the style
        grade_table = Table(table_data)
        grade_table.setStyle(table_style)

        # Add the table to the elements list
        elements.append(grade_table)
        elements.append(Spacer(1, 12))

        elements.append(Paragraph(user_tests.footer, styles['Normal']))

        doc.build(elements)
        return response

    # Create the PDF document for teacher view
    return create_student_view_pdf(user_tests)


@login_required(login_url='/authentication/login')
def download_teacher_view_pdf(request, id):
    user_tests =UserTest.objects.get(owner=request.user, pk=id)
    alphabet = string.ascii_uppercase

    # Function to create the PDF document for teacher view
    def create_teacher_view_pdf(user_tests):
        response = HttpResponse(content_type='application/pdf')
        filename = f'{user_tests.header}_teacher_view.pdf'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'

        doc = SimpleDocTemplate(response, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []

        header_style = styles['Title']
        header_style.alignment = 1  # Center alignment
        elements.append(Paragraph(user_tests.header, header_style))

        subtitle_style = styles['Heading3']
        subtitle_style.alignment = 1  # Center alignment
        elements.append(Paragraph(user_tests.subtitle, subtitle_style))

        institution_style = styles['Heading3']
        institution_style.alignment = 1  # Center alignment
        elements.append(Paragraph(user_tests.institution, institution_style))

        add_header_info_style = styles['Heading3']
        add_header_info_style.alignment = 1  # Center alignment
        elements.append(Paragraph(user_tests.add_header_info, add_header_info_style))

        for question_id, question_data in user_tests.questions.items():
            elements.append(Paragraph(f"Question {question_id[1:]}:", styles['Heading2']))
            elements.append(Paragraph(question_data['question'], styles['Normal']))
            elements.append(Spacer(1, 12))

            if 'answers' in question_data:
                for index, answer in enumerate(question_data['answers'], start=1):
                    elements.append(Paragraph(f"{str(alphabet[index-1])}. {answer}", styles['Normal']))

            if 'correct_answer' in question_data:
                elements.append(Paragraph("Correct Answer:", styles['Heading4']))
                correct_answer_list = question_data['correct_answer']
                elements.append(Paragraph(", ".join(str(alphabet[answer-1]) for answer in correct_answer_list), styles['Normal']))

            if question_data['explanation']:
                elements.append(Paragraph("Explanation:", styles['Heading4']))
                elements.append(Paragraph(question_data['explanation'], styles['Normal']))

        elements.append(Paragraph("Grades", styles['Heading2']))
        # Add table with grades
        table_data = [['Percentage', 'Grade']]
        for grade in user_tests.grades:
            table_data.append([f"{grade['percentage']}%", grade['grade']])

        table_style = TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                  ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                  ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                  ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                  ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                  ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                  ('GRID', (0, 0), (-1, -1), 1, colors.black)])

        # Create the table and apply the style
        grade_table = Table(table_data)
        grade_table.setStyle(table_style)

        # Add the table to the elements list
        elements.append(grade_table)
        elements.append(Spacer(1, 12))

        elements.append(Paragraph(user_tests.footer, styles['Normal']))

        doc.build(elements)
        return response

    # Create the PDF document for teacher view
    return create_teacher_view_pdf(user_tests)


# -----------download pdf views---
@login_required(login_url='/authentication/login')
def download_teacher_view(request, id):
    user_tests =UserTest.objects.get(owner=request.user, pk=id)
    alphabet = string.ascii_uppercase

    # Function to create the Word document for teacher view
    def create_teacher_view_docx(user_tests):
        doc = docx.Document()
        heading = doc.add_heading(user_tests.header, level=1)
        heading.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading(user_tests.subtitle, level=3)
        subtitle.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        institution = doc.add_heading(user_tests.institution, level=3)
        institution.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        add_header_info = doc.add_heading(user_tests.add_header_info, level=3)
        add_header_info.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        for question_id, question_data in user_tests.questions.items():
            doc.add_heading(f"Question {question_id[1:]}:", level=2)
            doc.add_paragraph(question_data['question'])

            if 'answers' in question_data:
                for index, answer in enumerate(question_data['answers'], start=1):
                    doc.add_paragraph(f"{alphabet[index-1]}. {answer}")

            if 'correct_answer' in question_data:
                doc.add_heading("Correct Answer:", level=3)
                correct_answer_list = question_data['correct_answer']  # Assuming correct_answer is a list
                doc.add_paragraph(", ".join(str(alphabet[answer-1]) for answer in correct_answer_list))

            if question_data['explanation']:
                doc.add_heading("Explanation:", level=3)
                doc.add_paragraph(question_data['explanation'])

            doc.add_paragraph("")  # Adding a blank line between questions

        doc.add_heading(f"Grading", level=2)
        # Add a table with the grades
        grades_table = doc.add_table(rows=1, cols=2)
        grades_table.autofit = True

        # Add table header
        header_cells = grades_table.rows[0].cells
        header_cells[0].text = 'Percentage'
        header_cells[1].text = 'Grade'

        # Add data rows to the table
        for grade_info in user_tests.grades:
            row_cells = grades_table.add_row().cells
            row_cells[0].text = f"{grade_info['percentage']}%"
            row_cells[1].text = grade_info['grade']

        doc.add_paragraph("")

        doc.add_paragraph(user_tests.footer)

        return doc

    # Create the Word document for teacher view
    doc = create_teacher_view_docx(user_tests)

    # Create a response with the Word document as a downloadable file
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={user_tests.header}_teacher_view.docx'

    # Save the document to the response
    doc.save(response)

    return response


@login_required(login_url='/authentication/login')
def download_student_view(request, id):
    user_tests =UserTest.objects.get(owner=request.user, pk=id)
    alphabet = string.ascii_uppercase

    # Function to create the Word document
    def create_student_view_docx(user_tests):
        doc = docx.Document()
        heading = doc.add_heading(user_tests.header, level=1)
        heading.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading(user_tests.subtitle, level=3)
        subtitle.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        institution = doc.add_heading(user_tests.institution, level=3)
        institution.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        add_header_info = doc.add_heading(user_tests.add_header_info, level=3)
        add_header_info.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        for question_id, question_data in user_tests.questions.items():
            doc.add_heading(f"Question {question_id[1:]}:", level=2)
            doc.add_paragraph(question_data['question'])

            if 'answers' in question_data:
                for index, answer in enumerate(question_data['answers'], start=1):
                    doc.add_paragraph(f"{alphabet[index-1]}. {answer}")

            doc.add_paragraph("")  # Adding a blank line between questions

        doc.add_heading(f"Grading", level=2)
        # Add a table with the grades
        grades_table = doc.add_table(rows=1, cols=2)
        grades_table.autofit = True

        # Add table header
        header_cells = grades_table.rows[0].cells
        header_cells[0].text = 'Percentage'
        header_cells[1].text = 'Grade'

        # Add data rows to the table
        for grade_info in user_tests.grades:
            row_cells = grades_table.add_row().cells
            row_cells[0].text = f"{grade_info['percentage']}%"
            row_cells[1].text = grade_info['grade']

        doc.add_paragraph("")

        doc.add_paragraph(user_tests.footer)

        return doc

    # Create the Word document
    doc = create_student_view_docx(user_tests)

    # Create a response with the Word document as a downloadable file
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={user_tests.header}_student_view.docx'

    # Save the document to the response
    doc.save(response)

    return response


# -----------------END DOWNLOAD VIEWS----------------------


@login_required(login_url='/authentication/login')
def search_tests(request):
    if request.method == 'POST':
        search_str = json.loads(request.body).get('searchText')
        user_tests = UserTest.objects.filter(
            header__icontains=search_str, owner=request.user) | UserTest.objects.filter(
            notes__icontains=search_str, owner=request.user) | UserTest.objects.filter(
            created_at__icontains=search_str, owner=request.user)
        data = user_tests.values()
        return JsonResponse(list(data), safe=False)


# Create your views here.
@login_required(login_url='/authentication/login')
def my_tests(request):
    user_membership = UserMembership.objects.filter(user=request.user).first()

    user_tests = UserTest.objects.filter(owner=request.user)

    # Get sorting parameters from query parameters
    sort_column = request.GET.get('sort_column', 'created_at')
    sort_order = request.GET.get('sort_order', 'desc')  # Default to descending order

    # Apply sorting to the queryset
    if sort_order == 'desc':
        user_tests = user_tests.order_by('-' + sort_column)
    else:
        user_tests = user_tests.order_by(sort_column)

    paginator = Paginator(user_tests, 5)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    context = {
        'user_tests': user_tests,
        'page_obj': page_obj,
        'sort_column': sort_column,
        'sort_order': sort_order,
        'user_membership': user_membership.membership,
    }
    return render(request, 'mytests/index.html', context)


@login_required(login_url='/authentication/login')
def home_view(request, id=None):
    user_membership = UserMembership.objects.filter(user=request.user).first()
    try:
        user_tests = UserTest.objects.get(owner=request.user, pk=id)
    except:
        user_tests = UserTest.objects.filter(owner=request.user).latest('pk')

    alphabet = string.ascii_uppercase
    for key, question in user_tests.questions.items():
        if 'answers' in question:
            updated_answers = [f"{alphabet[i]}. {answer}" for i, answer in enumerate(question['answers'])]
            question['answers'] = updated_answers

        if 'correct_answer' in question:
            updated_correct_answer = [alphabet[i-1] for i in question['correct_answer']]
            question['correct_answer'] = updated_correct_answer

    context = {
        'user_tests': user_tests,
        'values': user_tests,
        'user_membership': user_membership.membership,
        'alphabet': alphabet
    }
    if request.method == 'GET':
        return render(request, 'mytests/home-view.html', context)
    if request.method == 'POST':
        header = request.POST['header']


@login_required(login_url='/authentication/login')
def edit_pt(request, id):
    user_tests = UserTest.objects.get(owner=request.user, pk=id)
    context = {
        'user_tests': user_tests,
        'values': user_tests,
    }
    if request.method == 'GET':
        return render(request, 'mytests/edit_pt.html', context)
    if request.method == 'POST':
        header = request.POST['header']

        if not header:
            messages.error(request, 'Header is required')
            return render(request, 'mytests/edit_pt.html', context)

        subtitle = request.POST['subtitle']
        institution = request.POST['institution']
        add_header_info = request.POST['add_header_info']

        tag = request.POST['notes']
        footer = request.POST['footer']
        grades = []
        for i in range(1, 20):  # Assuming there are four grades (you can adjust the range based on your actual data)
            grade_key = f'grade_{i}_grade'
            percentage_key = f'grade_{i}_percentage'
            grade = request.POST.get(grade_key)
            percentage = request.POST.get(percentage_key)

            if grade and percentage:
                grades.append({
                    'grade': grade,
                    'percentage': int(percentage),  # Convert percentage to an integer if required
                })

        # Update main fields
        user_tests.header = header
        user_tests.subtitle = subtitle
        user_tests.institution = institution
        user_tests.add_header_info = add_header_info
        user_tests.grades = grades
        user_tests.notes = tag
        user_tests.footer = footer

        questions = {}
        for key, value in request.POST.items():

            if key.startswith('questions_'):
                parts = key.split('_')
                question_id = parts[1]
                field_name = parts[2]

                if question_id not in questions:
                    questions[question_id] = {}

                if field_name == 'answers':
                    if 'answers' not in questions[question_id]:
                        questions[question_id]['answers'] = []
                    questions[question_id]['answers'].append(value)
                elif field_name == 'correct':
                    questions[question_id]["correct_answer"] = [int(s) for s in value.split(',')]
                else:
                    questions[question_id][field_name] = value
        user_tests.questions = questions
        user_tests.save()
        messages.success(request, f'Practice test {header} updated successfully')

        return redirect('my-tests')


def delete_pt(request, id):
    user_test = UserTest.objects.get(owner=request.user, pk=id)
    header = user_test.header
    user_test.delete()
    messages.success(request, f'Test {header} removed')
    return redirect('my-tests')
